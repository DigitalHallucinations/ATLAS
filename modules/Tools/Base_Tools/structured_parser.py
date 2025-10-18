"""Structured document parser for common office formats.

This module exposes :class:`StructuredParser`, a small utility that normalizes
text, tables, and lightweight metadata from a handful of popular document
formats.  It relies on optional third-party packages to keep the runtime light:

* ``pdfminer.six`` for PDF extraction,
* ``python-docx`` for Microsoft Word documents, and
* ``pytesseract`` with Pillow for OCR on images.

When a dependency is unavailable the parser raises a descriptive
``StructuredParserDependencyError`` so callers can surface the requirement to
users.  Every call enforces a configurable byte budget to prevent extremely
large payloads from overwhelming the tool.  The default limit is 5 MiB and can
be overridden via the ``ATLAS_STRUCTURED_PARSER_MAX_BYTES`` environment
variable.

All helpers return plain Python data structures (strings, lists, dictionaries)
so they can be serialized directly by the tool router.
"""

from __future__ import annotations

from dataclasses import dataclass, field
import csv
import io
import os
from pathlib import Path
from typing import Any, Iterable, Optional, Sequence
from html.parser import HTMLParser

__all__ = [
    "StructuredParser",
    "StructuredParserError",
    "StructuredParserDependencyError",
    "DocumentTooLargeError",
    "UnsupportedFormatError",
    "InvalidInputError",
    "StructuredDocument",
    "TableExtraction",
]


_ENV_MAX_BYTES = "ATLAS_STRUCTURED_PARSER_MAX_BYTES"
_DEFAULT_MAX_BYTES = 5 * 1024 * 1024  # 5 MiB


class StructuredParserError(RuntimeError):
    """Base class for structured parser failures."""


class StructuredParserDependencyError(StructuredParserError):
    """Raised when an optional third-party dependency is missing."""


class DocumentTooLargeError(StructuredParserError):
    """Raised when an input payload exceeds the configured byte limit."""


class UnsupportedFormatError(StructuredParserError):
    """Raised when the caller requests an unsupported document format."""


class InvalidInputError(StructuredParserError):
    """Raised when the caller provides an invalid combination of arguments."""


@dataclass(slots=True)
class TableExtraction:
    """Simple representation of a tabular structure."""

    rows: list[list[str]] = field(default_factory=list)
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass(slots=True)
class StructuredDocument:
    """Container for normalized document output."""

    text: Optional[str] = None
    tables: list[TableExtraction] = field(default_factory=list)
    metadata: dict[str, Any] = field(default_factory=dict)


class StructuredParser:
    """Parse multiple document formats into structured primitives."""

    _IMAGE_EXTENSIONS = {"png", "jpg", "jpeg", "tiff", "bmp", "gif"}

    def __init__(self, *, max_bytes: Optional[int] = None) -> None:
        if max_bytes is None:
            env_limit = os.environ.get(_ENV_MAX_BYTES)
            if env_limit:
                try:
                    max_bytes = int(env_limit)
                except ValueError as exc:  # pragma: no cover - defensive guard
                    raise DocumentTooLargeError(
                        f"Environment variable {_ENV_MAX_BYTES} must be an integer."
                    ) from exc
            else:
                max_bytes = _DEFAULT_MAX_BYTES

        if max_bytes <= 0:
            raise DocumentTooLargeError("Maximum byte limit must be positive.")

        self._max_bytes = max_bytes

    # Public API -----------------------------------------------------
    def parse(
        self,
        *,
        content: Optional[bytes] = None,
        path: Optional[str] = None,
        source_format: Optional[str] = None,
        target_formats: Optional[Sequence[str]] = None,
    ) -> dict[str, Any]:
        """Parse a document located at ``path`` or supplied as ``content``.

        Parameters
        ----------
        content:
            Raw document bytes.  When provided the caller must also pass
            ``source_format`` to identify the file type.
        path:
            Filesystem path to the document.  The parser will read the file into
            memory up to ``max_bytes`` and infer the format from the file
            extension when ``source_format`` is omitted.
        source_format:
            Optional explicit format hint.  Accepted values are ``pdf``,
            ``docx``, ``html``, ``csv``, or any of the supported image
            extensions.
        target_formats:
            Iterable of output components to include.  Valid entries are
            ``text``, ``tables``, and ``metadata``.  When omitted all components
            are returned.
        """

        if content is None and path is None:
            raise InvalidInputError("Either 'content' or 'path' must be provided.")

        if content is not None and path is not None:
            raise InvalidInputError("Provide only one of 'content' or 'path'.")

        if content is not None and source_format is None:
            raise InvalidInputError("'source_format' is required when passing raw bytes.")

        if path:
            content = self._read_file(path)
            if source_format is None:
                source_format = self._infer_format_from_path(path)

        assert content is not None  # for mypy - guarded above

        normalized_format = self._normalize_format(source_format)

        document = self._dispatch_parse(normalized_format, content)

        requested = self._normalize_targets(target_formats)
        output: dict[str, Any] = {}

        if "text" in requested:
            output["text"] = document.text
        if "tables" in requested:
            output["tables"] = [
                {"rows": table.rows, "metadata": table.metadata}
                for table in document.tables
            ]
        if "metadata" in requested:
            base_metadata = dict(document.metadata)
            base_metadata.setdefault("size_bytes", len(content))
            base_metadata.setdefault("format", normalized_format)
            base_metadata.setdefault("table_count", len(document.tables))
            output["metadata"] = base_metadata

        return output

    # Internal helpers -----------------------------------------------
    def _read_file(self, path: str) -> bytes:
        candidate = Path(path)
        if not candidate.exists() or not candidate.is_file():
            raise InvalidInputError(f"Path '{path}' does not exist or is not a file.")

        size = candidate.stat().st_size
        if size > self._max_bytes:
            raise DocumentTooLargeError(
                f"Document is {size} bytes which exceeds the limit of {self._max_bytes}."
            )

        data = candidate.read_bytes()
        if len(data) > self._max_bytes:
            raise DocumentTooLargeError(
                f"Document is {len(data)} bytes which exceeds the limit of {self._max_bytes}."
            )

        return data

    def _infer_format_from_path(self, path: str) -> str:
        suffix = Path(path).suffix.lower().lstrip(".")
        if not suffix:
            raise UnsupportedFormatError("Unable to infer document format from path.")
        return suffix

    def _normalize_format(self, source_format: Optional[str]) -> str:
        if not source_format:
            raise UnsupportedFormatError("Document format was not provided.")

        lowered = source_format.lower()
        if lowered in {"pdf", "docx", "html", "htm", "csv"}:
            return "html" if lowered == "htm" else lowered
        if lowered in self._IMAGE_EXTENSIONS or lowered == "image":
            return "image"

        raise UnsupportedFormatError(f"Unsupported document format '{source_format}'.")

    def _normalize_targets(self, targets: Optional[Sequence[str]]) -> tuple[str, ...]:
        valid = {"text", "tables", "metadata"}
        if not targets:
            return tuple(valid)

        normalized: list[str] = []
        for target in targets:
            candidate = str(target).lower().strip()
            if candidate not in valid:
                raise InvalidInputError(f"Unknown target format '{target}'.")
            if candidate not in normalized:
                normalized.append(candidate)
        return tuple(normalized)

    def _dispatch_parse(self, fmt: str, content: bytes) -> StructuredDocument:
        if len(content) > self._max_bytes:
            raise DocumentTooLargeError(
                f"Document is {len(content)} bytes which exceeds the limit of {self._max_bytes}."
            )

        if fmt == "pdf":
            return self._parse_pdf(content)
        if fmt == "docx":
            return self._parse_docx(content)
        if fmt == "html":
            return self._parse_html(content)
        if fmt == "csv":
            return self._parse_csv(content)
        if fmt == "image":
            return self._parse_image(content)

        raise UnsupportedFormatError(f"Unsupported document format '{fmt}'.")

    # Format handlers -------------------------------------------------
    def _parse_pdf(self, content: bytes) -> StructuredDocument:
        try:
            from pdfminer.high_level import extract_text
        except Exception as exc:  # pragma: no cover - dependency error path
            raise StructuredParserDependencyError(
                "pdfminer.six is required to parse PDF documents."
            ) from exc

        text_stream = io.BytesIO(content)
        try:
            text = extract_text(text_stream)
        except Exception as exc:  # pragma: no cover - defensive
            raise StructuredParserError("Failed to parse PDF document.") from exc

        metadata = {
            "extracted_with": "pdfminer.six",
        }

        return StructuredDocument(text=text.strip() or None, metadata=metadata)

    def _parse_docx(self, content: bytes) -> StructuredDocument:
        try:
            from docx import Document  # type: ignore
        except Exception as exc:  # pragma: no cover - dependency error path
            raise StructuredParserDependencyError(
                "python-docx is required to parse DOCX documents."
            ) from exc

        try:
            document = Document(io.BytesIO(content))
        except Exception as exc:  # pragma: no cover - defensive
            raise StructuredParserError("Failed to parse DOCX document.") from exc

        paragraphs = [para.text.strip() for para in document.paragraphs if para.text.strip()]
        text = "\n".join(paragraphs) if paragraphs else None

        tables: list[TableExtraction] = []
        for index, table in enumerate(document.tables):
            rows: list[list[str]] = []
            for row in table.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            tables.append(
                TableExtraction(
                    rows=rows,
                    metadata={"index": index, "cell_count": sum(len(r) for r in rows)},
                )
            )

        metadata = {
            "extracted_with": "python-docx",
            "paragraph_count": len(paragraphs),
        }

        return StructuredDocument(text=text, tables=tables, metadata=metadata)

    def _parse_html(self, content: bytes) -> StructuredDocument:
        text: Optional[str] = None
        tables: list[TableExtraction] = []
        metadata: dict[str, Any] = {}

        decoded = content.decode("utf-8", errors="replace")

        try:
            from bs4 import BeautifulSoup  # type: ignore
        except Exception:
            # Fallback to a very small HTML parser when BeautifulSoup is not
            # installed.  This keeps the tool functional for plain text
            # extraction even without the optional dependency.

            class _FallbackHTMLParser(HTMLParser):
                def __init__(self) -> None:
                    super().__init__(convert_charrefs=True)
                    self._chunks: list[str] = []

                def handle_data(self, data: str) -> None:
                    cleaned = data.strip()
                    if cleaned:
                        self._chunks.append(cleaned)

                def get_text(self) -> str:
                    return " ".join(self._chunks)

            fallback = _FallbackHTMLParser()
            fallback.feed(decoded)
            text = fallback.get_text() or None
            metadata["extracted_with"] = "html.parser"
            return StructuredDocument(text=text, metadata=metadata)

        soup = BeautifulSoup(decoded, "html.parser")
        text = soup.get_text(separator=" ", strip=True) or None

        title = soup.title.string.strip() if soup.title and soup.title.string else None

        for index, table in enumerate(soup.find_all("table")):
            rows: list[list[str]] = []
            for row in table.find_all("tr"):
                cells = row.find_all(["th", "td"])
                if not cells:
                    continue
                rows.append(
                    [cell.get_text(separator=" ", strip=True) for cell in cells]
                )
            if rows:
                tables.append(
                    TableExtraction(
                        rows=rows,
                        metadata={
                            "index": index,
                            "header": rows[0] if rows else [],
                        },
                    )
                )

        metadata = {
            "title": title,
            "table_count": len(tables),
            "extracted_with": "beautifulsoup4",
        }

        return StructuredDocument(text=text, tables=tables, metadata=metadata)

    def _parse_csv(self, content: bytes) -> StructuredDocument:
        decoded = content.decode("utf-8", errors="replace")
        reader = csv.reader(io.StringIO(decoded))
        rows = [list(row) for row in reader]

        tables = [
            TableExtraction(
                rows=rows,
                metadata={
                    "header": rows[0] if rows else [],
                    "row_count": len(rows),
                },
            )
        ]

        text_lines = [", ".join(row) for row in rows]
        text = "\n".join(text_lines) if text_lines else None

        metadata = {
            "extracted_with": "csv",
            "row_count": len(rows),
        }

        return StructuredDocument(text=text, tables=tables, metadata=metadata)

    def _parse_image(self, content: bytes) -> StructuredDocument:
        try:
            from PIL import Image
        except Exception as exc:  # pragma: no cover - dependency error path
            raise StructuredParserDependencyError(
                "Pillow is required to decode image inputs before OCR."
            ) from exc

        try:
            import pytesseract
        except Exception as exc:  # pragma: no cover - dependency error path
            raise StructuredParserDependencyError(
                "pytesseract is required to perform OCR on images."
            ) from exc

        try:
            with Image.open(io.BytesIO(content)) as image:
                text = pytesseract.image_to_string(image) or None
                width, height = image.size
                mode = image.mode
        except StructuredParserError:
            raise
        except Exception as exc:  # pragma: no cover - defensive
            raise StructuredParserError("Failed to OCR image input.") from exc

        metadata = {
            "extracted_with": "pytesseract",
            "width": width,
            "height": height,
            "mode": mode,
        }

        return StructuredDocument(text=text, metadata=metadata)
